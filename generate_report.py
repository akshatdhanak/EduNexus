"""
Generate EduNexus Project Report in Word format.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    # Set background shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F3F4F6')
    run._element.rPr.append(shading)
    return p


# =====================================================================
# TITLE PAGE
# =====================================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EduNexus')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Unified Academic Utility System')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Project Report')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Technology Stack: ').bold = True
info.add_run('Django 4.2 · Python 3.13 · PostgreSQL · Bootstrap 5 · Google Gemini AI')

doc.add_paragraph()

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
info2.add_run('Deployment: ').bold = True
info2.add_run('Render (https://edunexus-hgae.onrender.com)')

doc.add_paragraph()
info3 = doc.add_paragraph()
info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
info3.add_run('GitHub: ').bold = True
info3.add_run('https://github.com/akshatdhanak/EduNexus')

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (manual)
# =====================================================================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Technology Stack',
    '3. Project Folder Structure',
    '4. Database Schema (24 Models)',
    '5. Application Modules & Features',
    '   5.1 Authentication & Role-Based Access',
    '   5.2 Admin Module',
    '   5.3 Faculty Module',
    '   5.4 Student Module',
    '6. Key Algorithms & Formulas',
    '   6.1 Attendance Percentage',
    '   6.2 Attendance Risk Predictor',
    '   6.3 Plagiarism Detection (Jaccard Similarity)',
    '   6.4 Timetable Conflict Detection',
    '   6.5 Grading System',
    '   6.6 Fee Calculation',
    '7. AI Database Chat (Gemini Integration)',
    '8. API Endpoints',
    '9. Environment Variables',
    '10. Installation & Setup Guide',
    '11. Deployment Configuration',
    '12. Screenshots',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =====================================================================
# 1. PROJECT OVERVIEW
# =====================================================================
add_heading_styled('1. Project Overview', level=1)
doc.add_paragraph(
    'EduNexus is a comprehensive, full-stack university management system designed to unify '
    'academic operations — attendance tracking, examination management, timetable scheduling, '
    'fee management, assignment workflows, and AI-powered analytics — into a single, '
    'role-based web application.'
)
doc.add_paragraph(
    'The system serves three primary user roles — Admin, Faculty, and Student — each with '
    'dedicated dashboards, feature sets, and access controls. The application is built with '
    'Django (Python) on the backend, Bootstrap 5 on the frontend, and is deployed on Render '
    'with a managed PostgreSQL database.'
)

add_heading_styled('Key Highlights', level=2)
highlights = [
    'Role-based access control (Admin, Faculty, Student)',
    'Real-time attendance tracking with barcode scanning support',
    'Automated attendance risk prediction with mathematical modeling',
    'Plagiarism detection using Jaccard Similarity for assignment submissions',
    'Timetable conflict detection using interval overlap algorithm',
    'AI-powered database chat using Google Gemini API (Admin only)',
    'Comprehensive exam management with admit cards and result processing',
    'Fee structure management with auto-calculation of outstanding amounts',
    'Leave request workflow with approval/rejection system',
    'Notification system for cross-role communication',
    'Deployed on Render with PostgreSQL, WhiteNoise static files, and Gunicorn',
]
for h in highlights:
    add_bullet(h)

doc.add_page_break()

# =====================================================================
# 2. TECHNOLOGY STACK
# =====================================================================
add_heading_styled('2. Technology Stack', level=1)

add_table(
    ['Category', 'Technology', 'Version / Details'],
    [
        ['Backend Framework', 'Django', '4.2.28'],
        ['Programming Language', 'Python', '3.13'],
        ['Database (Production)', 'PostgreSQL', '16 (Render Managed)'],
        ['Database (Development)', 'SQLite3', 'Built-in'],
        ['Frontend', 'Bootstrap 5', '5.3.0'],
        ['Icons', 'Material Icons + Bootstrap Icons', 'Latest'],
        ['Fonts', 'Google Fonts (Inter)', 'Latest'],
        ['Charts', 'Chart.js', '4.4.0'],
        ['AI Integration', 'Google Gemini API', 'gemini-2.0-flash'],
        ['Barcode Scanning', 'OpenCV + pyzbar', '4.13 / 0.1.9'],
        ['Image Processing', 'Pillow', '12.1.1'],
        ['WSGI Server', 'Gunicorn', '23.0.0'],
        ['Static Files', 'WhiteNoise', '6.8.2'],
        ['DB Connection', 'dj-database-url', '2.3.0'],
        ['Deployment Platform', 'Render', 'Free Tier'],
        ['Version Control', 'Git + GitHub', '—'],
    ]
)

doc.add_page_break()

# =====================================================================
# 3. PROJECT FOLDER STRUCTURE
# =====================================================================
add_heading_styled('3. Project Folder Structure', level=1)
doc.add_paragraph('The project follows Django\'s standard app-based architecture:')

structure = """EduNexus/
├── manage.py                  # Django management entry point
├── requirements.txt           # Python dependencies
├── render.yaml                # Render deployment blueprint
├── Dockerfile                 # Docker production image
├── docker-compose.yml         # Local Docker Compose setup
├── build.sh                   # Render build script
├── entrypoint.sh              # Docker entrypoint script
├── db.sqlite3                 # Local SQLite database
│
├── project1/                  # Django project settings
│   ├── settings.py            # Main settings (DB, middleware, etc.)
│   ├── urls.py                # Root URL configuration
│   ├── wsgi.py                # WSGI application
│   ├── middleware.py           # Custom session & role middleware
│   └── context_processors.py  # Template context processors
│
├── registration/              # Authentication app
│   ├── models.py              # CustomUser model (AbstractUser)
│   ├── views.py               # Login, logout, role redirect
│   ├── forms.py               # Login form
│   └── templates/registration/
│
├── admin_app/                 # Admin module (core models + views)
│   ├── models.py              # 24 database models
│   ├── views.py               # Admin CRUD operations
│   ├── chat_views.py          # AI Database Chat (Gemini)
│   ├── forms.py               # Student & Faculty forms
│   ├── exam_views.py          # Exam management views
│   ├── signals.py             # Auto-enrollment signals
│   ├── management/commands/   # Custom management commands
│   └── templates/admin_app/
│
├── faculty_app/               # Faculty module
│   ├── views.py               # Attendance, assignments, plagiarism
│   └── templates/faculty_app/
│
├── student_app/               # Student module
│   ├── views.py               # Dashboard, risk predictor, submissions
│   └── templates/student_app/
│
├── templates/                 # Shared templates
│   └── layout.html            # Base layout with nav
│
├── static/                    # Static assets (images)
├── staticfiles/               # Collected static files (WhiteNoise)
└── media/                     # User uploads (photos)"""

add_code_block(structure)

doc.add_page_break()

# =====================================================================
# 4. DATABASE SCHEMA
# =====================================================================
add_heading_styled('4. Database Schema (24 Models)', level=1)
doc.add_paragraph(
    'The system uses a normalized relational database with 24 core entities. '
    'Below is the complete list of models organized by domain:'
)

add_heading_styled('Academic Structure', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['Department', 'Academic departments', 'name, code, head (FK→Faculty)'],
        ['DegreeProgram', 'Degree programs offered', 'name, code, department (FK), duration_semesters'],
        ['Subject', 'Academic subjects', 'name, code, semester, credits, department (FK)'],
        ['SubjectOffering', 'Subject instances per division', 'subject (FK), faculty (FK), division, academic_year'],
    ]
)

add_heading_styled('User Entities', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['CustomUser', 'Auth user (AbstractUser)', 'username, email, password, role (admin/faculty/student)'],
        ['Student', 'Student profile', 'user (1:1), roll_number, degree_program (FK), semester, division, batch'],
        ['Faculty', 'Faculty profile', 'user (1:1), name, department, salary, subjects (M2M)'],
        ['StudentEnrollment', 'Student-subject mapping', 'student (FK), subject_offering (FK), status, enrollment_date'],
    ]
)

add_heading_styled('Examination Entities', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['ExamType', 'Types of exams', 'name (midterm/final/practical/quiz)'],
        ['ExamSchedule', 'Exam schedule entries', 'subject (FK), exam_type (FK), date, time, total_marks'],
        ['AdmitCard', 'Student admit cards', 'student (FK), exam_schedule (FK), is_approved'],
        ['ExamMarks', 'Marks obtained', 'student (FK), exam_schedule (FK), marks_obtained, is_marked'],
        ['InternalAssessment', 'Internal assessment marks', 'student (FK), subject_offering (FK), marks, max_marks'],
    ]
)

add_heading_styled('Attendance & Timetable', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['Lecture', 'Individual lecture sessions', 'subject_offering (FK), date, start_time, end_time, topic'],
        ['Attendance', 'Attendance records', 'student (FK), lecture (FK), status (present/absent/late)'],
        ['Timetable', 'Weekly timetable slots', 'subject_offering (FK), day, start_time, end_time, room_number'],
    ]
)

add_heading_styled('Results & Fees', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['SemesterResult', 'Semester-wise results', 'student (FK), semester, academic_year, sgpa, status'],
        ['SubjectResult', 'Subject-wise results', 'semester_result (FK), subject (FK), total_marks, grade, gpa'],
        ['FeeStructure', 'Fee records', 'student (FK), semester, fees_to_be_collected, paid, outstanding'],
        ['FeeReceipt', 'Payment receipts', 'fee_structure (FK), amount, payment_mode, transaction_id'],
    ]
)

add_heading_styled('Communication & Assignments', level=2)
add_table(
    ['Model', 'Description', 'Key Fields'],
    [
        ['LeaveRequest', 'Leave applications', 'user (FK), leave_type, start_date, end_date, status'],
        ['Notification', 'System notifications', 'title, message, sender (FK), recipient_role'],
        ['AcademicCalendar', 'Calendar events', 'title, event_type, start_date, end_date'],
        ['Assignment', 'Faculty assignments', 'title, subject (FK), created_by (FK), due_date, max_marks'],
        ['AssignmentSubmission', 'Student submissions', 'assignment (FK), student (FK), content, plagiarism_score'],
    ]
)

doc.add_page_break()

# =====================================================================
# 5. APPLICATION MODULES & FEATURES
# =====================================================================
add_heading_styled('5. Application Modules & Features', level=1)

# 5.1
add_heading_styled('5.1 Authentication & Role-Based Access', level=2)
doc.add_paragraph(
    'The system uses Django\'s built-in authentication extended with a CustomUser model '
    '(AbstractUser) that adds a "role" field with choices: admin, faculty, student. '
    'Three custom middleware classes enforce access control:'
)
add_bullet('SessionTimeoutMiddleware', bold_prefix='')
add_bullet(' — 30-minute sliding session window', bold_prefix='SessionTimeoutMiddleware')
add_bullet(' — Only one active session per user', bold_prefix='SingleSessionMiddleware')
add_bullet(' — URL-level role restriction (admin_app → admin only, etc.)', bold_prefix='RoleBasedAccessMiddleware')

# 5.2
add_heading_styled('5.2 Admin Module', level=2)
admin_features = [
    ('User Management — ', 'CRUD operations for Students and Faculty with form validation, '
     'username uniqueness checks, and auto-enrollment on semester changes.'),
    ('Attendance Management — ', 'Mark faculty attendance, view attendance reports.'),
    ('Timetable Management — ', 'Create, edit, delete timetable entries. Conflict detection for overlapping slots.'),
    ('Subject Management — ', 'CRUD for subjects with department assignment and semester mapping.'),
    ('Exam Management — ', 'Create exam schedules, manage admit cards, enter marks, publish results.'),
    ('Leave Management — ', 'Approve/reject leave requests from students and faculty.'),
    ('Notification System — ', 'Send notifications to specific roles or all users.'),
    ('AI Database Chat — ', 'Natural language queries to the database using Google Gemini API. '
     'Generates SQL, executes safely, returns tables and auto-generated charts.'),
    ('Timetable Conflict Optimizer — ', 'Detects overlapping timetable slots using interval overlap algorithm.'),
]
for prefix, desc in admin_features:
    add_bullet(desc, bold_prefix=prefix)

# 5.3
add_heading_styled('5.3 Faculty Module', level=2)
faculty_features = [
    ('Mark Attendance — ', 'Select subject offering, create lecture, mark students present/absent/late.'),
    ('View Attendance Reports — ', 'Subject-wise and date-wise attendance analytics.'),
    ('Assignment Management — ', 'Create assignments for subjects, set due dates and max marks.'),
    ('Plagiarism Checker — ', 'Run Jaccard Similarity-based plagiarism detection on all submissions '
     'for an assignment. Flags pairs above 70% similarity with common phrase extraction.'),
    ('Profile Management — ', 'Edit personal details and view profile.'),
    ('View Timetable — ', 'See personal teaching schedule.'),
    ('Leave Requests — ', 'Apply for leave with date range and reason.'),
]
for prefix, desc in faculty_features:
    add_bullet(desc, bold_prefix=prefix)

# 5.4
add_heading_styled('5.4 Student Module', level=2)
student_features = [
    ('View Attendance — ', 'Subject-wise attendance with percentage breakdown (theory + practical).'),
    ('Attendance Risk Predictor — ', 'AI-driven risk analysis per subject. Predicts future attendance '
     'if classes are skipped. Calculates recovery classes needed. Shows safe-skip count.'),
    ('View Timetable — ', 'Personal weekly timetable view.'),
    ('View Results — ', 'Semester-wise and subject-wise academic results with grades and SGPA.'),
    ('View Fees — ', 'Fee structure, payment history, outstanding calculations.'),
    ('Submit Assignments — ', 'Submit text content for faculty-created assignments.'),
    ('Timetable Conflicts — ', 'View detected scheduling conflicts in personal timetable.'),
    ('Profile Management — ', 'Edit personal details, change password.'),
    ('Leave Requests — ', 'Apply for leave, view status of past requests.'),
    ('Notifications — ', 'View system-wide and role-specific notifications.'),
    ('Exam Admit Cards — ', 'View and download admit cards for upcoming exams.'),
]
for prefix, desc in student_features:
    add_bullet(desc, bold_prefix=prefix)

doc.add_page_break()

# =====================================================================
# 6. KEY ALGORITHMS & FORMULAS
# =====================================================================
add_heading_styled('6. Key Algorithms & Formulas', level=1)

# 6.1
add_heading_styled('6.1 Attendance Percentage', level=2)
doc.add_paragraph('Calculated per-subject or overall for a student:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Attendance % = (Classes Attended / Total Classes) × 100')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph('Minimum required attendance: 75%')

# 6.2
add_heading_styled('6.2 Attendance Risk Predictor', level=2)
doc.add_paragraph('The risk predictor uses the following calculations:')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Risk Level Classification:').bold = True
add_table(
    ['Risk Level', 'Condition', 'Color'],
    [
        ['Safe', 'Attendance ≥ 80%', 'Green'],
        ['Warning', '75% ≤ Attendance < 80%', 'Yellow'],
        ['Critical', 'Attendance < 75%', 'Red'],
    ]
)

p = doc.add_paragraph()
p.add_run('Classes Required to Recover (reach 75%):').bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('Required = ⌈0.75 × Total Classes⌉ − Attended')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Future Prediction (if student skips N more classes):').bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('Future % = (Attended / (Total + N)) × 100')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Maximum Classes Student Can Still Skip:').bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('Can Skip = ⌊Attended / 0.75⌋ − Total Classes')
run.bold = True
run.font.size = Pt(11)

# 6.3
add_heading_styled('6.3 Plagiarism Detection (Jaccard Similarity)', level=2)
doc.add_paragraph('The plagiarism checker uses a three-step approach:')

p = doc.add_paragraph()
p.add_run('Step 1 — Text Preprocessing:').bold = True
doc.add_paragraph(
    'Convert to lowercase → Remove all punctuation → Split into words → '
    'Remove 80+ English stop words (the, a, is, are, etc.) → Discard single-character words'
)

p = doc.add_paragraph()
p.add_run('Step 2 — Jaccard Similarity (pairwise between all submissions):').bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('J(A, B) = |A ∩ B| / |A ∪ B| × 100')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph(
    'Where A and B are the sets of preprocessed words from two submissions. '
    'The result is a percentage (0–100%).'
)

p = doc.add_paragraph()
p.add_run('Step 3 — Common Phrase Detection (N-gram Matching):').bold = True
phrases_steps = [
    'Generate n-grams from both texts (minimum length = 4 words)',
    'Find exact n-gram matches between the two submissions',
    'Deduplicate: remove phrases that are substrings of longer matches',
    'Return top 10 longest matched phrases',
]
for s in phrases_steps:
    add_bullet(s)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Flagging Threshold: ').bold = True
p.add_run('Similarity ≥ 70% → Flagged as plagiarized')

# 6.4
add_heading_styled('6.4 Timetable Conflict Detection', level=2)
doc.add_paragraph('For each day, timetable slots are grouped by division. Two slots conflict if:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Slot_A.start < Slot_B.end  AND  Slot_B.start < Slot_A.end')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph(
    'This is the standard interval overlap algorithm. '
    'The system checks all slot pairs within the same division and reports conflicts '
    'with subject names, faculty, room numbers, and time ranges.'
)

# 6.5
add_heading_styled('6.5 Grading System', level=2)
doc.add_paragraph('Grades and GPA are auto-calculated based on total marks:')
add_table(
    ['Grade', 'Marks Range', 'GPA Points'],
    [
        ['A+', '90 – 100', '4.0'],
        ['A', '80 – 89', '3.7'],
        ['B+', '70 – 79', '3.3'],
        ['B', '60 – 69', '3.0'],
        ['C', '50 – 59', '2.0'],
        ['F', 'Below 50', '0.0'],
    ]
)
doc.add_paragraph('Total Marks = Internal Marks + External Marks + Practical Marks')
doc.add_paragraph('Pass/Fail: Any grade except F = Pass')

# 6.6
add_heading_styled('6.6 Fee Calculation', level=2)
doc.add_paragraph('Outstanding fees are auto-calculated on every save:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Outstanding = Fees to be Collected − Paid − Refunded')
run.bold = True
run.font.size = Pt(11)

doc.add_page_break()

# =====================================================================
# 7. AI DATABASE CHAT
# =====================================================================
add_heading_styled('7. AI Database Chat (Gemini Integration)', level=1)
doc.add_paragraph(
    'The AI Database Chat allows administrators to query the database using natural language. '
    'It uses Google\'s Gemini API (gemini-2.0-flash model) to translate questions into SQL.'
)

p = doc.add_paragraph()
p.add_run('How It Works:').bold = True
steps = [
    'Admin types a question like "Show students with attendance below 60%"',
    'The system sends the question along with the complete database schema to Gemini AI',
    'Gemini generates a safe, read-only SQL query (SELECT only)',
    'The system executes the query against the database',
    'Results are formatted as HTML tables in the chat interface',
    'Gemini also suggests appropriate Chart.js visualizations (bar, pie, line charts)',
    'Charts are rendered dynamically using Chart.js 4.4.0',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'{step}')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Security: ').bold = True
p.add_run(
    'Only SELECT queries are allowed. INSERT, UPDATE, DELETE, DROP, and other '
    'write operations are blocked. Only admin users can access this feature.'
)

doc.add_page_break()

# =====================================================================
# 8. API ENDPOINTS
# =====================================================================
add_heading_styled('8. API Endpoints (URL Routes)', level=1)

add_heading_styled('Authentication', level=2)
add_table(
    ['URL', 'Method', 'Description'],
    [
        ['/accounts/login/', 'GET/POST', 'Login page with role selection'],
        ['/accounts/logout/', 'POST', 'Logout and clear session'],
        ['/health/', 'GET', 'Health check endpoint (returns JSON)'],
    ]
)

add_heading_styled('Admin Endpoints', level=2)
add_table(
    ['URL', 'Method', 'Description'],
    [
        ['/admin_app/admin_dashboard/', 'GET', 'Admin dashboard'],
        ['/admin_app/.../student_add', 'GET/POST', 'Register new student'],
        ['/admin_app/.../<id>/student_edit', 'GET/POST', 'Edit student'],
        ['/admin_app/.../<id>/student_delete', 'POST', 'Delete student'],
        ['/admin_app/.../faculty_add', 'GET/POST', 'Register new faculty'],
        ['/admin_app/.../<id>/faculty_edit', 'GET/POST', 'Edit faculty'],
        ['/admin_app/.../manage_timetable', 'GET', 'View timetable'],
        ['/admin_app/.../manage_subjects/', 'GET', 'Subject management'],
        ['/admin_app/.../notification', 'GET/POST', 'Send notifications'],
        ['/admin_app/.../leave', 'GET', 'Manage leave requests'],
        ['/admin_app/.../timetable_conflicts/', 'GET', 'Detect conflicts'],
        ['/admin_app/.../database_chat/', 'GET', 'AI Chat interface'],
        ['/admin_app/.../chat_api/', 'POST', 'AI Chat API (AJAX)'],
    ]
)

add_heading_styled('Faculty Endpoints', level=2)
add_table(
    ['URL', 'Method', 'Description'],
    [
        ['/faculty_app/faculty_dashboard/', 'GET', 'Faculty dashboard'],
        ['/faculty_app/mark_attendance/', 'GET/POST', 'Mark attendance'],
        ['/faculty_app/view_attendance/', 'GET', 'View attendance reports'],
        ['/faculty_app/assignments/', 'GET', 'List assignments'],
        ['/faculty_app/assignments/create/', 'GET/POST', 'Create assignment'],
        ['/faculty_app/assignments/<id>/plagiarism/', 'GET', 'Run plagiarism check'],
    ]
)

add_heading_styled('Student Endpoints', level=2)
add_table(
    ['URL', 'Method', 'Description'],
    [
        ['/student_app/student_dashboard/', 'GET', 'Student dashboard'],
        ['/student_app/view_attendance/', 'GET', 'View attendance'],
        ['/student_app/view_timetable/', 'GET', 'View timetable'],
        ['/student_app/view_results/', 'GET', 'View results'],
        ['/student_app/view_fees/', 'GET', 'View fees'],
        ['/student_app/assignments/', 'GET', 'View assignments'],
        ['/student_app/assignments/<id>/submit/', 'POST', 'Submit assignment'],
        ['/student_app/attendance_risk/', 'GET', 'Risk predictor'],
        ['/student_app/timetable_conflicts/', 'GET', 'View conflicts'],
    ]
)

doc.add_page_break()

# =====================================================================
# 9. ENVIRONMENT VARIABLES
# =====================================================================
add_heading_styled('9. Environment Variables', level=1)
doc.add_paragraph(
    'The following environment variables must be set for production deployment. '
    'A .env.example file should be created with placeholder values:'
)

add_table(
    ['Variable', 'Description', 'Example'],
    [
        ['SECRET_KEY', 'Django secret key', '<auto-generated>'],
        ['DEBUG', '0 for production, 1 for development', '0'],
        ['DATABASE_URL', 'PostgreSQL connection string', 'postgres://user:pass@host:5432/db'],
        ['ALLOWED_HOSTS', 'Comma-separated hostnames', 'edunexus-hgae.onrender.com'],
        ['CSRF_TRUSTED_ORIGINS', 'Comma-separated HTTPS origins', 'https://edunexus-hgae.onrender.com'],
        ['GEMINI_API_KEY', 'Google Gemini API key', 'AIzaSy...'],
        ['DJANGO_SUPERUSER_USERNAME', 'Initial admin username', 'admin'],
        ['DJANGO_SUPERUSER_PASSWORD', 'Initial admin password', '<set in dashboard>'],
        ['DJANGO_SUPERUSER_EMAIL', 'Initial admin email', 'admin@example.com'],
        ['PYTHON_VERSION', 'Python version for Render', '3.13.0'],
    ]
)

doc.add_page_break()

# =====================================================================
# 10. INSTALLATION & SETUP GUIDE
# =====================================================================
add_heading_styled('10. Installation & Setup Guide', level=1)

add_heading_styled('Prerequisites', level=2)
add_bullet('Python 3.13+')
add_bullet('pip (Python package manager)')
add_bullet('PostgreSQL (for production) or SQLite (for development)')
add_bullet('Git')

add_heading_styled('Step 1: Clone the Repository', level=2)
add_code_block('git clone https://github.com/akshatdhanak/EduNexus.git\ncd EduNexus')

add_heading_styled('Step 2: Create Virtual Environment', level=2)
add_code_block('python -m venv .venv\nsource .venv/bin/activate    # macOS/Linux\n.venv\\Scripts\\activate       # Windows')

add_heading_styled('Step 3: Install Dependencies', level=2)
add_code_block('pip install -r requirements.txt')

add_heading_styled('Step 4: Set Environment Variables', level=2)
doc.add_paragraph('Create a .env file (or export variables) based on .env.example:')
add_code_block('export SECRET_KEY="your-secret-key"\nexport DEBUG=1\nexport GEMINI_API_KEY="your-gemini-api-key"')

add_heading_styled('Step 5: Run Migrations', level=2)
add_code_block('python manage.py migrate')

add_heading_styled('Step 6: Seed Default Data', level=2)
add_code_block('python manage.py seed_departments')

add_heading_styled('Step 7: Create Superuser', level=2)
add_code_block('python manage.py createsuperuser')

add_heading_styled('Step 8: Collect Static Files', level=2)
add_code_block('python manage.py collectstatic --noinput')

add_heading_styled('Step 9: Run Development Server', level=2)
add_code_block('python manage.py runserver')
doc.add_paragraph('Open http://127.0.0.1:8000/ in your browser.')

doc.add_page_break()

# =====================================================================
# 11. DEPLOYMENT CONFIGURATION
# =====================================================================
add_heading_styled('11. Deployment Configuration', level=1)

add_heading_styled('Render Deployment (render.yaml)', level=2)
doc.add_paragraph('The project includes a render.yaml blueprint for one-click deployment on Render:')
deploy_features = [
    'Web service with Python runtime (free tier)',
    'Managed PostgreSQL database (free tier)',
    'Auto-deploy on GitHub push',
    'Build command: ./build.sh (installs deps + collectstatic)',
    'Start command: migrate → seed_departments → create_superuser_if_missing → gunicorn',
    'Health check at /health/',
    'Environment variables auto-configured',
]
for f in deploy_features:
    add_bullet(f)

add_heading_styled('Docker Deployment', level=2)
doc.add_paragraph('For containerized deployment:')
add_code_block('docker-compose up --build')
doc.add_paragraph(
    'This starts a PostgreSQL 16 container and the Django app with Gunicorn. '
    'Media files are persisted via Docker volumes.'
)

add_heading_styled('Production Security Features', level=2)
security = [
    'HTTPS with SSL redirect (SECURE_SSL_REDIRECT)',
    'HSTS enabled (31,536,000 seconds / 1 year)',
    'Secure cookies (CSRF_COOKIE_SECURE, SESSION_COOKIE_SECURE)',
    'Content type nosniff (SECURE_CONTENT_TYPE_NOSNIFF)',
    'XSS filter (SECURE_BROWSER_XSS_FILTER)',
    'X-Frame-Options: DENY',
    'Proxy SSL header support for Render\'s reverse proxy',
]
for s in security:
    add_bullet(s)

doc.add_page_break()

# =====================================================================
# 12. SCREENSHOTS
# =====================================================================
add_heading_styled('12. Screenshots', level=1)
doc.add_paragraph(
    '(Insert screenshots of the following pages here:)'
)
screenshots = [
    'Login Page',
    'Admin Dashboard',
    'Faculty Dashboard',
    'Student Dashboard',
    'Student Registration Form',
    'Mark Attendance Page',
    'Attendance Risk Predictor',
    'Plagiarism Check Results',
    'AI Database Chat',
    'Timetable Conflict Detection',
    'Exam Results Page',
    'Fee Structure Page',
]
for s in screenshots:
    add_bullet(s)
    # Add placeholder box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Screenshot Placeholder]')
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'EduNexus_Project_Report.docx')
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
